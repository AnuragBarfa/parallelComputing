import os
import csv
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt

files=['naiveOutput.txt','recursiveOutput.txt','reduceOutput.txt']
data=[]
for file in files:
	with open(file,'r') as csvfile:
		cursor = csv.reader(csvfile, delimiter=' ')
		temp_data=[]
		for row in cursor:
			temp_data.append(row)
		data.append(temp_data)	
cleanedData=[[] for _ in range(len(files)+1) ]
print(cleanedData)
for fileIndex in range(0,len(data)):
	rowIndex=0
	while rowIndex<len(data[fileIndex]):
		if fileIndex==0:
			cleanedData[0].append(float(data[fileIndex][rowIndex][2]))
		cleanedData[fileIndex+1].append(float(data[fileIndex][rowIndex+1][4]))
		rowIndex=rowIndex+2
print(cleanedData)
plt.figure()
plt.xlabel('Matrix size (n)')
plt.ylabel('Computation time(ns)')
plt.title('Graph showing variation of time for addition as size changes')
plt.plot(cleanedData[0],cleanedData[1],label='naive')
plt.plot(cleanedData[0],cleanedData[2],label='using recursion')
plt.plot(cleanedData[0],cleanedData[3],label='using reduce')
# plt.plot(n,t1,label='optimization1')
# plt.plot(n,t2,label='optimization2')
# plt.plot(n,t3,label='optimization3')
plt.legend(loc='upper left')
plt.savefig('comparisionGraph.png')
# plt.show()
try:
	os.remove('report.docx')
except OSError:
	pass

document = Document()

document.add_heading('Parallel Computing Assignment', 0)

p = document.add_paragraph('')
p.add_run('Roll No. - ').bold = True
p.add_run(' 160123005 ')

table = document.add_table(rows=1, cols=len(files)+1)

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Size(n)'
for i in range(0,len(files)):
	hdr_cells[i+1].text=files[i].split(".")[0]

for i in range(0,len(cleanedData[0])):
    row_cells = table.add_row().cells
    row_cells[0].text = str(cleanedData[0][i])
    for j in range(0,len(files)):
    	row_cells[j+1].text=str(cleanedData[j+1][i])
r=p.add_run()
r.add_picture('comparisionGraph.png')
document.add_page_break()

document.save('report.docx')			

# p.add_run('italic.').italic = True

# document.add_heading('Heading, level 1', level=1)

# document.add_paragraph(
#     'first item in unordered list', style='List Bullet'
# )
# document.add_paragraph(
#     'first item in ordered list', style='List Number'
# )

# document.add_picture('monty-truth.png', width=Inches(1.25))
